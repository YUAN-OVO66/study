#比第二版多引入了  docx2python 库，用于将表名文本和表对象联系起来， 以期将标题嵌入表格内容，提高检索准确性
#尝试将文本标题嵌入下级内容中，提高检索准确性
#将表对象对应的表面在向量数据库中，单独存储为一项字段，以供进行专项处理
import os
import datetime
from pymilvus import connections, Collection, utility, db, MilvusClient
#from pymilvus.model.hybrid import BGEM3EmbeddingFunction
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
from modelscope.outputs import OutputKeys
from modelscope.pipelines import pipeline
from modelscope.utils.constant import Tasks
from docx import Document
from docx2markdown import docx_to_markdown
from docx2python import docx2python
from convert_table_to_markdown import convert_table_to_markdown, convert_table_to_markdown_without_head, convert_table_to_markdown_each_row
from text_chunking_funcation import check_if_title

doc_path= "C:\\Users\\LV\\Desktop\\IMC\\UOH_AIService\\syllabus\\农业机械化.docx"
doc = Document(doc_path)  # 打开文档
lines =[]

# 单独处理纯文本
str = ""
for p in range(len(doc.paragraphs)):
    doc.paragraphs[p].text =  doc.paragraphs[p].text.strip()
    if  doc.paragraphs[p].text != "":
        if  doc.paragraphs[p].text.endswith("。") or  doc.paragraphs[p].text.endswith("！") or  doc.paragraphs[p].text.endswith("？") or   doc.paragraphs[p].text.endswith("!") or   doc.paragraphs[p].text.endswith("?")  or   doc.paragraphs[p].text.endswith("：")  or   doc.paragraphs[p].text.endswith(":") or   doc.paragraphs[p].text.endswith("；")  or   doc.paragraphs[p].text.endswith(";"):
            str = str+ doc.paragraphs[p].text 
        else:
            if len( doc.paragraphs[p].text) > 30:
                str = str +  doc.paragraphs[p].text+"。 " 
            else:
                if check_if_title( doc.paragraphs[p].text):
                    if p != 0:
                        if check_if_title( doc.paragraphs[p-1].text):
                                str = str +  doc.paragraphs[p].text+"： " 
                                continue
                    print( doc.paragraphs[p].text)
                    str = str +" \r\n "+  doc.paragraphs[p].text+"： " 
                else:
                    str = str +  doc.paragraphs[p].text+"： " 

p = pipeline(
    task=Tasks.document_segmentation,
    model='damo/nlp_bert_document-segmentation_chinese-base')

result = p(documents= str);
lines = [line for line in result[OutputKeys.TEXT].splitlines() if line.strip() != ""]

#对于过长的文本，再切一次
l=0
while l < len(lines) :
    if len(lines[l]) > 1000:
        result_part = p(documents= lines[l]);
        lines_new = [line for line in result_part[OutputKeys.TEXT].splitlines() if line.strip() != ""]
        if  len(lines_new) > 1:
            lines[l+1:l+1] = lines_new
            # lines.extend(lines_new)
        else:
            split_point = (len(lines[l]) + 1) // 2
            lines[l+1:l+1]  = [lines[l][:split_point], lines[l][split_point:]]
            # lines.extend([lines[l][:split_point], lines[l][split_point:]])
        lines.remove(lines[l])
        l=l-1
    l = l+1

#记录下纯文本部分的切块数
lines_paragraph_count = len(lines)  

# print(lines)

 #docx2python读取文本测试
doc2 = docx2python(doc_path)  # 打开文档
# print(doc2.body)
table_title = []
for t in range(len(doc.tables)):  
# for t in range(2,3):  
    for i in range(len(doc2.body)):
        first_text =""
        for x in range(len(doc2.body[i][0][0])):
            first_text = first_text+ doc2.body[i][0][0][x].strip()
        if doc.tables[t].rows[0].cells[0].text.replace("\n","").replace("\r","").strip() ==   first_text:
            # print(1)
            got = False
            if len(doc2.body[i][0]) >2 : 
                # print(2, doc.tables[t].rows[0].cells[1].text.strip() , doc2.body[i][0][1][0].strip())
                # print(2.1, doc.tables[t].rows[0].cells[2].text.strip() , doc2.body[i][0][2][0].strip())
                if   doc.tables[t].rows[0].cells[1].text.strip() == doc2.body[i][0][1][0].strip() and doc.tables[t].rows[0].cells[2].text.strip() == doc2.body[i][0][2][0].strip():
                    # print(2.5)
                    got = True
                # 处理表头中有换行符的情况
                elif len(doc2.body[i][0][1]) >1 and len(doc2.body[i][0][2]) > 1:       
                    if doc.tables[t].rows[0].cells[1].text.strip().replace("\n","").replace("\r","") == doc2.body[i][0][1][0].strip() + doc2.body[i][0][1][1].strip() and doc.tables[t].rows[0].cells[2].text.strip().replace("\n","").replace("\r","") == doc2.body[i][0][2][0].strip() + doc2.body[i][0][2][1].strip():
                        # print(2.5)
                        got = True
            elif len(doc2.body[i]) >1 : 
                # print(3)
                if   doc.tables[t].rows[1].cells[0].text.strip() == doc2.body[i][1][0][0].strip()  :
                    # print(3.5)
                    got = True
            if got:
                # print(4)
                title__txt = doc2.body[i-1][len(doc2.body[i-1])-1][len(doc2.body[i-1][len(doc2.body[i-1])-1])-1][ len(doc2.body[i-1][len(doc2.body[i-1])-1][len(doc2.body[i-1][len(doc2.body[i-1])-1])-1])-1].strip()
                if ( title__txt.endswith("。") or  title__txt.endswith(".") or title__txt.endswith("；") or title__txt.endswith(";") or title__txt.endswith("？") or title__txt.endswith("?") or title__txt.endswith("！") or title__txt.endswith("!") ) and ( len(doc2.body[i-1][len(doc2.body[i-1])-1][len(doc2.body[i-1][len(doc2.body[i-1])-1])-1]) > 1):
                    title__txt =  doc2.body[i-1][len(doc2.body[i-1])-1][len(doc2.body[i-1][len(doc2.body[i-1])-1])-1][ len(doc2.body[i-1][len(doc2.body[i-1])-1][len(doc2.body[i-1][len(doc2.body[i-1])-1])-1])-2].strip() + title__txt
                table_title.append( title__txt)
                break
        if i ==len(doc2.body) -1:
               table_title.append("")                                                                                                                                                                                                                                                               
    
# print(table_title)
table_title_update = []

# 单独处理表格
for tc in range(len(doc.tables)):  # 遍历所有表格
    has_head = True
    cell_span1= []
    cell_span2= []
    iF_fill2 = False
    
    # 判断是否应该把该表作为有表头的表进行处理
    if len(doc.tables[tc].rows) < 3:
        has_head = False
    else:
        for r in range( len(doc.tables[tc].rows)):
            c =0
            while c < len(doc.tables[tc].rows[r].cells):
                if  len(doc.tables[tc].rows[r].cells[c].text) > 50 and r ==0:
                    # print('>50')
                    has_head = False
                    break
                if len(cell_span1) != len(doc.tables[tc].rows[r].cells):
                    cell_span1.append(doc.tables[tc].rows[r].cells[c].grid_span)
                    continue
                if doc.tables[tc].rows[r].cells[c].grid_span != cell_span1[c]  and iF_fill2 == False:
                    iF_fill2 = True 
                    for cell_copy in doc.tables[tc].rows[r].cells:
                        cell_span2.append(cell_copy.grid_span)
                    break
                if iF_fill2:
                    if doc.tables[tc].rows[r].cells[c].grid_span != cell_span1[c] and doc.tables[tc].rows[r].cells[c].grid_span != cell_span2[c]:
                        has_head = False
                        break
                c=c+1
            if has_head == False:
                break
            

    if has_head:
        table_txt = convert_table_to_markdown(doc.tables[tc] )
        if len(table_txt) < 2000 :
            lines.append(table_title[tc] +  " \r\n   " +table_txt.replace("<br>", "") )  #同时嵌入表标题
            table_title_update.append(table_title[tc] )
        else:
            lines.extend(convert_table_to_markdown_each_row(table_title[tc] , doc.tables[tc], table_title_update)) #同时嵌入表标题
    else:           
        lines.append(table_title[tc] +  " \r\n   " +convert_table_to_markdown_without_head(doc.tables[tc]).replace("<br>", "")) #同时嵌入表标题
        table_title_update.append(table_title[tc] )


print(lines)
 
