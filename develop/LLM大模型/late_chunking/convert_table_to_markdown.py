from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_alignment_symbol(paragraph):
    """将 Word 段落对齐方式转换为 Markdown 语法"""
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: ':---',  # 左对齐
        WD_ALIGN_PARAGRAPH.CENTER: ':---:',  # 居中对齐
        WD_ALIGN_PARAGRAPH.RIGHT: '---:',  # 右对齐
        None: '---'  # 默认左对齐
    }
    return alignment_map.get(paragraph.alignment, '---')  # 返回对应的 Markdown 对齐符号，如果没有对应的符号，则返回默认的左对齐符号

def convert_table_to_markdown(table):
    """将单个 Word 表格转换为 Markdown 格式"""
    markdown = []
    
    # 处理表头
    headers = []
    alignments = []
    for cell in table.rows[0].cells:
        # 转义特殊字符并处理换行
        text = cell.text.replace('\n', '<br>').replace('|', r'\|')
        headers.append(text)
        
        # 获取对齐方式（取单元格第一个段落）
        if cell.paragraphs:
            para = cell.paragraphs[0]
            alignments.append(get_alignment_symbol(para))
        else:
            alignments.append('---')

    # 构建表头
    markdown.append(f"| {' | '.join(headers)} |")
    markdown.append(f"| {' | '.join(alignments)} |")

    # 处理数据行
    col_count = len(headers)
    for row in table.rows[1:]:
        cells = []
        for cell in row.cells:
            text = cell.text.replace('\n', '<br>').replace('|', r'\|')
            cells.append(text)
        
        # 处理列数不一致的情况（合并单元格可能导致列数变化）
        cells = cells[:col_count]  # 截断多余列
        while len(cells) < col_count:  # 补充缺失列
            cells.append('')
        
        markdown.append(f"| {' | '.join(cells)} |")

    return '\n'.join(markdown)

def convert_table_to_markdown_without_head(table):
    """将单个 Word 表格转换为五表头的 Markdown 格式"""
    markdown = []
    
    # 处理数据行
    col_count = len(table.rows[0].cells)
    for row in table.rows:
        cells = []
        # for cell in row.cells:
        #     text = cell.text.replace('\n', '<br>').replace('|', r'\|')
        #     cells.append(text)
        c = 0
        while c< len(row.cells):
            text = row.cells[c].text.replace('\n', '<br>').replace('|', r'\|')
            cells.append(text)
            c=c+row.cells[c].grid_span
        
        # 处理列数不一致的情况（合并单元格可能导致列数变化）
        cells = cells[:col_count]  # 截断多余列
        while len(cells) < col_count:  # 补充缺失列
            cells.append('')
        
        markdown.append(f"| {' | '.join(cells)} |")

    return '\n'.join(markdown)


def convert_table_to_markdown_each_row(table_title, table, table_title_update):
    """将单个 Word 表格分割为 单行带表头的多表 Markdown 格式"""
    markdown = []
    
    # 处理表头
    headers = []
    alignments = []
    for cell in table.rows[0].cells:
        # 转义特殊字符并处理换行
        text = cell.text.replace('\n', '<br>').replace('|', r'\|')
        headers.append(text)
        
        # 获取对齐方式（取单元格第一个段落）
        if cell.paragraphs:
            para = cell.paragraphs[0]
            alignments.append(get_alignment_symbol(para))
        else:
            alignments.append('---')

 
    # 处理数据行
    col_count = len(headers)
    for row in table.rows[1:]:
        row_tex = ""
        cells = []
        for cell in row.cells:
            text = cell.text.replace('\n', '<br>').replace('|', r'\|')
            cells.append(text)
        
        # 处理列数不一致的情况（合并单元格可能导致列数变化）
        cells = cells[:col_count]  # 截断多余列
        while len(cells) < col_count:  # 补充缺失列
            cells.append('')
        # 构建表头
        row_tex =  f"| {' | '.join(headers)} |"
        row_tex = row_tex + f"| {' | '.join(alignments)} |"
        row_tex = row_tex + f"| {' | '.join(cells)} |"
        markdown.append(table_title+  " \r\n   " +row_tex.replace("<br>", ""))
        table_title_update.append(table_title)
    return  markdown

